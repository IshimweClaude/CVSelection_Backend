from rest_framework import permissions, filters
from rest_framework.views import APIView
from rest_framework.generics import RetrieveUpdateDestroyAPIView
from rest_framework.response import Response
from .models import Country
from .serializers import *
from django.utils import timezone
from django.db import IntegrityError
# Record AFDB countries
from rest_framework import status
import json
from django_filters.rest_framework import DjangoFilterBackend
from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from authentication.models import Applicant
from django.core.exceptions import ObjectDoesNotExist
from authentication.models import User


class CountryCreateView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    
    def post(self, request):
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can add new countries."}, status=status.HTTP_403_FORBIDDEN)

        serializer = CountrySerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Country added successfully!", "data": serializer.data}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# List all countries
class CountryListView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['country_name','isMemberOfAFDB','id']
    search_fields = ['country_name','isMemberOfAFDB','id']

    def get(self, request):
        countries = Country.objects.all()
        if countries:
            serializer = CountrySerializer(countries, many=True)
            return Response({"message": "Countries fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        else:
            return Response({"message": "No countries found."}, status=status.HTTP_404_NOT_FOUND)
        
class CountryDetailView(RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Country.objects.all()
    serializer_class = CountrySerializer
    lookup_field = 'id'
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['country_name','isMemberOfAFDB','id']
    search_fields = ['country_name','isMemberOfAFDB','id']
    
    def get(self, request, id):
        try:
            country = Country.objects.get(id=id)
        except Country.DoesNotExist:
            return Response({"detail": "Country not found."}, status=status.HTTP_404_NOT_FOUND)
        serializer = CountrySerializer(country)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def put(self, request, id):
        try:
            country = Country.objects.get(id=id)
        except Country.DoesNotExist:
            return Response({"detail": "Country not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can update countries."}, status=status.HTTP_403_FORBIDDEN)
        serializer = CountrySerializer(country, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Country updated successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id):
        try:
            country = Country.objects.get(id=id)
        except Country.DoesNotExist:
            return Response({"detail": "Country not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can delete countries."}, status=status.HTTP_403_FORBIDDEN)
        country.delete()
        return Response({"message": "Country deleted successfully!"}, status=status.HTTP_200_OK)

# Record AFDB jobs
class JobCreateView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend,filters.SearchFilter,filters.OrderingFilter]
    filterset_fields = ['id','title','location','status','job_type','id','salary','postedDate','deadline']
    search_fields = ['id','title','location','status','job_type','id','salary','postedDate','deadline']

    def post(self, request):
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can add new jobs."}, status=status.HTTP_403_FORBIDDEN)

        serializer = JobSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Job added successfully!", "data": serializer.data}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

# List all jobs

class JobListView(APIView):
    permission_classes = [permissions.AllowAny]
    
    filter_backends = [DjangoFilterBackend,filters.SearchFilter,filters.OrderingFilter]
    filterset_fields = ['id','title','location','status','job_type','id','salary','postedDate','deadline']
    search_fields = ['id','title','location','status','job_type','id','salary','postedDate','deadline']
    authentication_classes = []
    
    def get(self, request):
        jobs = Job.objects.filter(status='open')
        
        if jobs:
            serializer = JobSerializer(jobs, many=True)
            return Response({"message": "Jobs fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        else:
            return Response({"message": "No jobs found."}, status=status.HTTP_404_NOT_FOUND)


class JobDetailView(RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Job.objects.all()
    serializer_class = JobSerializer
    lookup_field = 'id'
    filter_backends = [DjangoFilterBackend,filters.SearchFilter,filters.OrderingFilter]
    filterset_fields = ['id','title','location','status','job_type','id','salary','postedDate','deadline']
    search_fields = ['id','title','location','status','job_type','id','salary','postedDate','deadline']
        
    def get(self, request, id):
        try:
            job = Job.objects.get(id=id)
        except Job.DoesNotExist:
            return Response({"detail": "Job not found."}, status=status.HTTP_404_NOT_FOUND)
        serializer = JobSerializer(job)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def put(self, request, id):
        try:
            job = Job.objects.get(id=id)
        except Job.DoesNotExist:
            return Response({"detail": "Job not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can update jobs."}, status=status.HTTP_403_FORBIDDEN)
        serializer = JobSerializer(job, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Job updated successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id):
        try:
            job = Job.objects.get(id=id)
        except Job.DoesNotExist:
            return Response({"detail": "Job not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can delete jobs."}, status=status.HTTP_403_FORBIDDEN)
        job.delete()
        return Response({"message": "Job deleted successfully!"}, status=status.HTTP_200_OK)
    

# ====================
# Record AFDB formal education as part of applicant's profile

# class Formal_educationCreateView(APIView):
#     permission_classes = [permissions.IsAuthenticated]

#     def post(self, request):
#         if request.user.user_role != 'applicant':
#             return Response({"detail": "Sorry, only applicants can add formal education."}, status=status.HTTP_403_FORBIDDEN)

#         serializer = Formal_educationSerializer(data=request.data)
#         if serializer.is_valid():
#             serializer.save()
#             return Response({"message": "Formal education added successfully!", "data": serializer.data}, status=status.HTTP_201_CREATED)
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

# list my formal education

# class Formal_educationListView(APIView):
#     permission_classes = [permissions.IsAuthenticated]

#     def get(self, request):
#         formal_education = Formal_education.objects.filter(applicant=request.user)
#         if formal_education:
#             serializer = Formal_educationSerializer(formal_education, many=True)
#             return Response({"message": "Formal education fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
#         else:
#             return Response({"message": "No formal education found."}, status=status.HTTP_404_NOT_FOUND)
        
# Record AFDB work experience as part of applicant's profile

# class Work_experienceCreateView(APIView):
#     permission_classes = [permissions.IsAuthenticated]

#     def post(self, request):
#         if request.user.user_role != 'applicant':
#             return Response({"detail": "Sorry, only applicants can add work experience."}, status=status.HTTP_403_FORBIDDEN)

#         serializer = Work_experienceSerializer(data=request.data)
#         if serializer.is_valid():
#             serializer.save()
#             return Response({"message": "Work experience added successfully!", "data": serializer.data}, status=status.HTTP_201_CREATED)
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

# list my work experience

# class Work_experienceListView(APIView):
#     permission_classes = [permissions.IsAuthenticated]

#     def get(self, request):
#         work_experience = Work_experience.objects.filter(applicant=request.user)
#         if work_experience:
#             serializer = Work_experienceSerializer(work_experience, many=True)
#             return Response({"message": "Work experience fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
#         else:
#             return Response({"message": "No work experience found."}, status=status.HTTP_404_NOT_FOUND)
        

# Record AFDB language skills as part of applicant's profile

# class Language_skillsCreateView(APIView):
#     permission_classes = [permissions.IsAuthenticated]

#     def post(self, request):
#         if request.user.user_role != 'applicant':
#             return Response({"detail": "Sorry, only applicants can add language skills."}, status=status.HTTP_403_FORBIDDEN)

#         serializer = Language_skillsSerializer(data=request.data)
#         if serializer.is_valid():
#             serializer.save()
#             return Response({"message": "Language skills added successfully!", "data": serializer.data}, status=status.HTTP_201_CREATED)
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# list my language skills

# class Language_skillsListView(APIView):
#     permission_classes = [permissions.IsAuthenticated]

#     def get(self, request):
#         language_skills = Language_skills.objects.filter(applicant=request.user)
#         if language_skills:
#             serializer = Language_skillsSerializer(language_skills, many=True)
#             return Response({"message": "Language skills fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
#         else:
#             return Response({"message": "No language skills found."}, status=status.HTTP_404_NOT_FOUND)
        
class JobApplicationView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, job_id):
        # Check if the job exists and has not reached the deadline
        try:
            job = Job.objects.get(id=job_id)
        except Job.DoesNotExist:
            return Response({"detail": "Job not found."}, status=status.HTTP_404_NOT_FOUND)

        if job.deadline < timezone.now().date():
            return Response({"detail": "Job application deadline has passed."}, status=status.HTTP_403_FORBIDDEN)

        # Check if the user is an applicant
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicants can apply for jobs."}, status=status.HTTP_403_FORBIDDEN)

        # Create the application
        application_data = request.data.copy()
        application_data['applicant'] = request.user
        application_data['job'] = job.id
        application_serializer = ApplicationSerializer(data=application_data)
        if application_serializer.is_valid():
            try:
                application_serializer.save()
            except IntegrityError:
                return Response({"detail": "You have already applied for this job."}, status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response(application_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        # Create the formal education
        formal_educations_data = request.data.get('formal_educations')
        if formal_educations_data is not None:
            try:
                formal_educations_data = json.loads(formal_educations_data)
            except json.JSONDecodeError:
                return Response({"detail": "Invalid data. 'formal_educations' should be a JSON string."}, status=status.HTTP_400_BAD_REQUEST)

        if not isinstance(formal_educations_data, list):
            return Response({"detail": "Invalid data. 'formal_educations' should be a list."}, status=status.HTTP_400_BAD_REQUEST)

        for formal_education_data in formal_educations_data:
            formal_education_data['applicant'] = request.user.id
            try:
                country = Country.objects.get(country_name=formal_education_data['country'])
            except Country.DoesNotExist:
                return Response({"detail": "Country not found."}, status=status.HTTP_404_NOT_FOUND)
            formal_education_data['country'] = country.id
            formal_education_serializer = Formal_educationSerializer(data=formal_education_data)
            if formal_education_serializer.is_valid():
                formal_education_serializer.save()
            else:
                return Response(formal_education_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
       # Create the work experience
        work_experiences_data = request.data.get('work_experiences')
        if work_experiences_data is not None:
            try:
                work_experiences_data = json.loads(work_experiences_data)
            except json.JSONDecodeError:
                return Response({"detail": "Invalid data. 'work_experiences' should be a JSON string."}, status=status.HTTP_400_BAD_REQUEST)

        if not isinstance(work_experiences_data, list):
            return Response({"detail": "Invalid data. 'work_experiences' should be a list."}, status=status.HTTP_400_BAD_REQUEST)

        for work_experience_data in work_experiences_data:
            work_experience_data['applicant'] = request.user.id
            try:
                country = Country.objects.get(country_name=work_experience_data['working_country'])
            except Country.DoesNotExist:
                return Response({"detail": "Country not found."}, status=status.HTTP_404_NOT_FOUND)
            work_experience_data['working_country'] = country.id
            work_experience_serializer = Work_experienceSerializer(data=work_experience_data)
            if work_experience_serializer.is_valid():
                work_experience_serializer.save()
            else:
                return Response(work_experience_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        # Create the language skills
        language_skills_data = request.data.get('language_skills')
        if language_skills_data is not None:
            try:
                language_skills_data = json.loads(language_skills_data)
            except json.JSONDecodeError:
                return Response({"detail": "Invalid data. 'language_skills' should be a JSON string."}, status=status.HTTP_400_BAD_REQUEST)

        if not isinstance(language_skills_data, list):
            return Response({"detail": "Invalid data. 'language_skills' should be a list."}, status=status.HTTP_400_BAD_REQUEST)

        for language_skill_data in language_skills_data:
            language_skill_data['applicant'] = request.user.id
            language_skills_serializer = Language_skillsSerializer(data=language_skill_data)
            if language_skills_serializer.is_valid():
                language_skills_serializer.save()
            else:
                return Response(language_skills_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
        
        applicant_id = request.user.id
        applicant_data = request.data.get('applicant')

        try:
            applicant_data = json.loads(applicant_data)
        except json.JSONDecodeError:
            return Response({"detail": "Invalid data. 'applicant' should be a JSON string."}, status=status.HTTP_400_BAD_REQUEST)

        # Retrieve the existing applicant object
        try:
            applicant = Applicant.objects.get(applicant_id=applicant_id)
        except Applicant.DoesNotExist:
            return Response({"detail": "Applicant not found"}, status=status.HTTP_404_NOT_FOUND)
        
        # Retrieve the country object
        try:
            country = Country.objects.get(country_name=applicant_data.get('country'))
        except ObjectDoesNotExist:
            return Response({"detail": "Country not found"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Update the country in applicant_data
        applicant_data['country'] = country.id

        # Update the applicant object
        applicant_serializer = ApplicantSerializer(applicant, data=applicant_data, partial=True)
        if applicant_serializer.is_valid():
            applicant_serializer.save()
        else:
            return Response(applicant_serializer.errors, status=status.HTTP_400_BAD_REQUEST)   

        return Response({"message": "Applicant Info updated successfully!"}, status=status.HTTP_200_OK)
# ========================================================================================================
class JobApplicationListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can view job applications."}, status=status.HTTP_403_FORBIDDEN)

        applications = Application.objects.all()
        if applications:
            serializer = ApplicationSerializer(applications, many=True)
            return Response({"message": "Job applications fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        else:
            return Response({"message": "No job applications found."}, status=status.HTTP_404_NOT_FOUND)
        
class JobApplicationListViewById(APIView):
    permission_classes = [permissions.IsAuthenticated]
    def get(self, request, id):
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can view job applications."}, status=status.HTTP_403_FORBIDDEN)

        try:
            application = Application.objects.get(id=id)
        except Application.DoesNotExist:
            return Response({"detail": "Application not found."}, status=status.HTTP_404_NOT_FOUND)
        serializer = ApplicationSerializer(application)
        return Response(serializer.data, status=status.HTTP_200_OK)
# ===================================================

class JobApplicationDetailView(RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Application.objects.all()
    serializer_class = ApplicationSerializer
    lookup_field = 'id'
    
    def get(self, request, id):
        try:
            job_application = Application.objects.get(id=id)
        except Application.DoesNotExist:
            return Response({"detail": "Job application not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view job application."}, status=status.HTTP_403_FORBIDDEN)
        serializer = ApplicationSerializer(job_application)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def put(self, request, id):
        try:
            job_application = Application.objects.get(id=id)
        except Application.DoesNotExist:
            return Response({"detail": "Job application not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can update job application."}, status=status.HTTP_403_FORBIDDEN)
        serializer = ApplicationSerializer(job_application, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Job application updated successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id):
        try:
            job_application = Application.objects.get(id=id)
        except Application.DoesNotExist:
            return Response({"detail": "Job application not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can delete job application."}, status=status.HTTP_403_FORBIDDEN)
        job_application.delete()
        return Response({"message": "Job application deleted successfully!"}, status=status.HTTP_200_OK)



# Formal Education List
# ==========================
class FormalEducationListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view Language Skills."}, status=status.HTTP_403_FORBIDDEN)

        education = Formal_education.objects.filter(applicant=request.user.applicant)
        if education:
            serializer = Formal_educationSerializer(education, many=True)
            return Response({"message": "Formal Education data fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        else:
            return Response({"message": "No Formal Education data found."}, status=status.HTTP_404_NOT_FOUND)
            
# Formal Education Update and Delete
# ====================================
class FormalEducationDetailView(RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Formal_education.objects.all()
    serializer_class = Formal_educationSerializer
    lookup_field = 'id'
    
    def get(self, request, id):
        try:
            formal_education = Formal_education.objects.get(id=id)
        except Formal_education.DoesNotExist:
            return Response({"detail": "Formal education not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view formal education."}, status=status.HTTP_403_FORBIDDEN)
        serializer = Formal_educationSerializer(formal_education)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def put(self, request, id):
        try:
            formal_education = Formal_education.objects.get(id=id)
        except Formal_education.DoesNotExist:
            return Response({"detail": "Formal education not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can update formal education."}, status=status.HTTP_403_FORBIDDEN)
        serializer = Formal_educationSerializer(formal_education, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Formal education updated successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id):
        try:
            formal_education = Formal_education.objects.get(id=id)
        except Formal_education.DoesNotExist:
            return Response({"detail": "Formal education not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can delete formal education."}, status=status.HTTP_403_FORBIDDEN)
        formal_education.delete()
        return Response({"message": "Formal education deleted successfully!"}, status=status.HTTP_200_OK)
    
# # Work Experince List
# ====================================
class WorkExperienceListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view work experience."}, status=status.HTTP_403_FORBIDDEN)
        
        experience = Work_experience.objects.filter(applicant=request.user.applicant)
        if experience:
            serializer = Work_experienceSerializer(experience, many=True)
            return Response({"message": "Work Experience fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        else:
            return Response({"message": "No Work Experience found."}, status=status.HTTP_404_NOT_FOUND)
        
# Work Experince Update and Delete
# ====================================
class WorkExperienceDetailView(RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Work_experience.objects.all()
    serializer_class = Work_experienceSerializer
    lookup_field = 'id'
    
    def get(self, request, id):
        try:
            work_experience = Work_experience.objects.get(id=id)
        except Work_experience.DoesNotExist:
            return Response({"detail": "Work experience not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view work experience."}, status=status.HTTP_403_FORBIDDEN)
        serializer = Work_experienceSerializer(work_experience)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def put(self, request, id):
        try:
            work_experience = Work_experience.objects.get(id=id)
        except Work_experience.DoesNotExist:
            return Response({"detail": "Work experience not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can update work experience."}, status=status.HTTP_403_FORBIDDEN)
        serializer = Work_experienceSerializer(work_experience, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Work experience updated successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id):
        try:
            work_experience = Work_experience.objects.get(id=id)
        except Work_experience.DoesNotExist:
            return Response({"detail": "Work experience not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can delete work experience."}, status=status.HTTP_403_FORBIDDEN)
        work_experience.delete()
        return Response({"message": "Work experience deleted successfully!"}, status=status.HTTP_200_OK)

# # Language Skills List
# ==========================
class LanguageSkillsListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view Language Skills."}, status=status.HTTP_403_FORBIDDEN)

        language = Language_skills.objects.filter(applicant=request.user.applicant)
        if language:
            serializer = Language_skillsSerializer(language, many=True)
            return Response({"message": "Language Skills fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        else:
            return Response({"message": "No Language Skills found."}, status=status.HTTP_404_NOT_FOUND)
    
# # Language Skills update and Delete
# ===========================================
class LanguageSkillsDetailView(RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Language_skills.objects.all()
    serializer_class = Language_skillsSerializer
    lookup_field = 'id'
    
    def get(self, request, id):
        try:
            language_skills = Language_skills.objects.get(id=id)
        except Language_skills.DoesNotExist:
            return Response({"detail": "Language skills not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view language skills."}, status=status.HTTP_403_FORBIDDEN)
        serializer = Language_skillsSerializer(language_skills)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def put(self, request, id):
        try:
            language_skills = Language_skills.objects.get(id=id)
        except Language_skills.DoesNotExist:
            return Response({"detail": "Language skills not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can update language skills."}, status=status.HTTP_403_FORBIDDEN)
        serializer = Language_skillsSerializer(language_skills, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Language skills updated successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id):
        try:
            language_skills = Language_skills.objects.get(id=id)
        except Language_skills.DoesNotExist:
            return Response({"detail": "Language skills not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can delete language skills."}, status=status.HTTP_403_FORBIDDEN)
        language_skills.delete()
        return Response({"message": "Language skills deleted successfully!"}, status=status.HTTP_200_OK)
    

# Applicant List
# ==========================
class ApplicantListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view Language Skills."}, status=status.HTTP_403_FORBIDDEN)

        applicant = Applicant.objects.filter(applicant=request.user.applicant)
        if applicant:
            serializer = ApplicantSerializer(applicant, many=True)
            return Response({"message": "Applicant Information fetched successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        else:
            return Response({"message": "No Applicant Information found."}, status=status.HTTP_404_NOT_FOUND)
        
class ApplicantDetailView(RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Applicant.objects.all()
    serializer_class = ApplicantSerializer
    lookup_field = 'id'
    
    def get(self, request, id):
        try:
            applicant = Applicant.objects.get(id=id)
        except Applicant.DoesNotExist:
            return Response({"detail": "Applicant not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can view applicant profile."}, status=status.HTTP_403_FORBIDDEN)
        serializer = ApplicantSerializer(applicant)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def put(self, request, id):
        try:
            applicant = Applicant.objects.get(id=id)
        except Applicant.DoesNotExist:
            return Response({"detail": "Applicant not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can update applicant profile."}, status=status.HTTP_403_FORBIDDEN)
        serializer = ApplicantSerializer(applicant, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Applicant profile updated successfully!", "data": serializer.data}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, id):
        try:
            applicant = Applicant.objects.get(id=id)
        except Applicant.DoesNotExist:
            return Response({"detail": "Applicant not found."}, status=status.HTTP_404_NOT_FOUND)
        if request.user.user_role != 'applicant':
            return Response({"detail": "Sorry, only applicant can delete applicant profile."}, status=status.HTTP_403_FORBIDDEN)
        applicant.delete()
        return Response({"message": "Applicant profile deleted successfully!"}, status=status.HTTP_200_OK)
    
# ========================================================================================================
    # AI Integration part
# ========================================================================================================
class ProcessResumesAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, id):
        if request.user.user_role != 'recruiter':
            return Response({"detail": "Sorry, only recruiters can view job applications."}, status=status.HTTP_403_FORBIDDEN)
        
        try:
            job = Job.objects.get(id=id)
        except Job.DoesNotExist:
            return Response({"detail": "Job not found."}, status=status.HTTP_404_NOT_FOUND)
        try:
            applications = Application.objects.filter(job_id=id)
        except Application.DoesNotExist:
            return Response({"detail": "Job application not found."}, status=status.HTTP_404_NOT_FOUND)
        
        cv_dict = {}

        for application in applications:
            applicant_id = application.applicant_id
            cv = application.cv
            cv_dict[applicant_id] = str(cv.path)
        
                   
        engl_job_description = job.english_job_description.path
        fr_job_description = job.french_job_description.path

        # print("English Job Description: ", engl_job_description)
        # print("French Job Description: ", fr_job_description)
        # print("List of Applicant's cvs", cv_dict.keys, cv_dict.values)
        
        # Calling model function
        
        results = compare_job_descriptions_and_cvs (engl_job_description, fr_job_description,cv_dict.values())
        # results.sort(key=lambda x: x['score'], reverse=True) 
        # print(results)
        # results['score'] = results['score'].astype(float)
        # results = results.sort_values(by=['score'], ascending=False)
        for result in results:
            for cv_path in cv_dict.values():
                if cv_path == result['full_path']:
                    result['applicant_id'] = list(cv_dict.keys())[list(cv_dict.values()).index(cv_path)]
                    # result['score'] = result['score']
                    result['applicant_Name'] = User.objects.get(id=result['applicant_id']).first_name + " " + User.objects.get(id=result['applicant_id']).last_name
                    
        return Response(results, status=status.HTTP_200_OK)
        

# C:\Users\Claude Ishimwe\Documents\Docs\Other Skills\Python Django\AFDB\CVSelection_Backend\CVSelection_Backend\application_files\annette_cv_fr.docx
#===================================================================================================
# =============== AI PART ===========================================================================
#===================================================================================================

from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import LAParams, LTText, LTTextBox
from pdfminer.converter import PDFPageAggregator

def is_white_color_pdf(color):
    # Check if the color is white (255, 255, 255)
    return color == (1, 1, 1)  # PDF colors are in the range [0, 1]

def extract_non_white_text_from_pdf(file_path):
    # Extract text from PDF excluding white-colored text
    with open(file_path, 'rb') as file:
        parser = PDFParser(file)
        document = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        text_content = ''
        for page in PDFPage.create_pages(document):
            interpreter.process_page(page)
            layout = device.get_result()
            for element in layout:
                if isinstance(element, (LTText, LTTextBox)):
                    text = element.get_text().strip()
                    color = element.get_font().get_color()
                    if not is_white_color_pdf(color):
                        text_content += text + '\n'

    return text_content



from docx import Document
from docx.shared import RGBColor

def is_white_color_docx(color):
    # Check if the color is white (255, 255, 255)
    return color == RGBColor(255, 255, 255)

def extract_non_white_text_from_docx(file_path):
    # Extract text from Word document excluding white-colored text
    doc = Document(file_path)
    text_content = ''

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text.strip()
            color = run.font.color.rgb if run.font.color else RGBColor(0, 0, 0)  # Default to black if no color
            if not is_white_color_docx(color):
                text_content += text + ' '

    return text_content



# for convert a path function
import docx2txt
import textract
import json
from pdfminer.high_level import extract_text


# for convert paths function
import os

# for parse_json and ranking function
from langdetect import detect
import spacy
from time import gmtime, strftime
import pandas as pd
import numpy as np
import transformers
from sklearn.feature_extraction.text import TfidfVectorizer
from spacy.lang.fr.stop_words import STOP_WORDS as fr_stop
from spacy.lang.en.stop_words import STOP_WORDS as en_stop

import tqdm


def convert_file_to_json(file_path, fr=False):
    try:
        if file_path.endswith('.docx'):
            # Handle DOCX files
            text = docx2txt.process(file_path)
#             text = extract_non_white_text_from_docx(file_path)
        elif file_path.endswith('.pdf'):
            text = extract_text(file_path)
#             text =  extract_non_white_text_from_pdf(file_path)
                
        else:
            return None  # Unsupported file format

        # Create a JSON object with the "tdr" key and the extracted text
        json_data = [{"tdr": text}]

        # Serialize the JSON object to a JSON string with UTF-8 encoding
        json_string = json.dumps(json_data, ensure_ascii=False, indent=4).encode('utf-8').decode('utf-8')

        if fr:
          with open("tdr_fr.json", "w", encoding="utf-8") as json_file:
            json_file.write(json_string)
        else:
          # Write the JSON data to a file named "tdr.json"
          with open("tdr.json", "w", encoding="utf-8") as json_file:
              json_file.write(json_string)

        return "tdr.json"

    except Exception as e:
        return str(e)
    


def convert_paths_to_json(path_list, fr=False):
    result = []

    for file_path in path_list:
        try:
            if file_path.endswith('.docx'):
                # Handle DOCX files using docx2txt
                text = docx2txt.process(file_path)
            elif file_path.endswith('.pdf'):
                # Handle PDF files using textract with the 'pdfminer' method
                text = extract_text(file_path)
            else:
                text = "Unsupported file format"

            # Extract the file name (id.pdf) from the path
            file_name = os.path.basename(file_path)

            # Create a dictionary with the file name as the key and the extracted text as the value
            json_entry = {file_name: text}

            result.append(json_entry)

        except Exception as e:
            return str(e)

    try:
        # Serialize the result list to a JSON string with UTF-8 encoding
        json_string = json.dumps(result, ensure_ascii=False, indent=4).encode('utf-8').decode('utf-8')

        if fr:
          # Write the JSON data to a file named "CVs.json"
          with open("CVs_fr.json", "w", encoding="utf-8") as json_file:
              json_file.write(json_string)
        else:
          # Write the JSON data to a file named "CVs.json"
          with open("CVs.json", "w", encoding="utf-8") as json_file:
              json_file.write(json_string)

        return "CVs.json"

    except Exception as e:
        return str(e)
    
    



# final_stopwords_list = list(fr_stop) + list(en_stop) + ['[UNK]']

from string import digits
remove_digits = str.maketrans('', '', digits)

def parse_json(json_file):
    with open(json_file, encoding='utf-8') as json_file:
        contents = json.load(json_file)

    data = []
    file_names = []
    symbols = "!\"#$%&()*+-./:;<=>?@[\]^_`{|}~«»•"
    for c in contents:
        file_name = [k for k in c.keys()][0]
        file_names.append(file_name)
        sample = c[file_name]
        sample = sample.replace('\\n', '\n')
        sample = sample.replace('\\xef\\x83\\x98 ', '')
        sample = sample.replace('\\xef\\x82\\xa7 ', '')
        sample = sample.replace('\\xe2\\x80\\xa6', '')
        sample = sample.replace('\\xef\\x82\\xb7 ', '')
        sample = sample.replace('\\', '')
        sample = sample.replace('\\x0co ', ' ')
        sample = sample.replace('\xa0', ' ')
        sample = sample.replace('\t', '')
        sample = sample.replace(',', ' ')

        sample = sample.replace("'", "")
        for s in symbols:
          sample = sample.replace(s, " ")
        sentences = sample.lower().split('\n')
        text = ' '.join(sentences).translate(remove_digits)
        data.append(text)

    return data, file_names


def remove_stop_words(data, nlp):
    new_text = []
    for text in tqdm.tqdm(data):
        words = nlp(text)
        no_stopwords = ' '.join(t.text for t in words if not t.is_stop)
        new_text.append(' '.join(no_stopwords.split())) 
    return new_text


def ranking(cvs_json_fr, cvs_json_en, tdr_json, idx=0, output=None,seuil=False,nombreseuil=4000,ner=False,stop=False):



    print(strftime("%Y-%m-%d %H:%M:%S", gmtime()))


    tdrs, tdr_names = parse_json(tdr_json)
    LANG = detect(tdrs[0]).upper()
    print("langue : ",LANG)
    if LANG == 'FR':
        cvs, cvs_names  = parse_json(cvs_json_fr)

    else:
        cvs, cvs_names  = parse_json(cvs_json_en)

    tdr = [tdrs[idx]]

    if LANG == 'FR':
        nlp = spacy.load('fr_core_news_sm')
    else:
        nlp = spacy.load('en_core_web_sm')


    CVS=[]
    for cv in cvs:
        ch="\n".join([v for v in cv.split('\n') if v])
        ch2=" ".join([v for v in ch.split(' ') if v])
        CVS.append(ch2)
    cvs=CVS

    if stop:
        cvs = remove_stop_words(cvs, nlp)
        tdrs = remove_stop_words(tdrs, nlp)


    score_cos = pd.DataFrame(np.zeros((len(cvs), len(tdr))),
                             columns=['tdr%d'%i for i in range(len(tdr))],
                             index=cvs_names)
    print(strftime("%Y-%m-%d %H:%M:%S", gmtime()))
    for MODEL in ['BERT', 'DistilBERT', 'CamemBERT']:
        if LANG == 'FR':
            if MODEL == 'BERT':
                tokenizer = transformers.BertTokenizer.from_pretrained('bert-base-multilingual-cased')
            elif MODEL == 'DistilBERT':
                tokenizer = transformers.DistilBertTokenizer.from_pretrained('distilbert-base-multilingual-cased')
            elif MODEL == 'CamemBERT' and LANG == 'FR':
                tokenizer = transformers.CamembertTokenizer.from_pretrained('camembert-base')
        #elif LANG == 'EN':
        else:
            if MODEL == 'BERT':
                tokenizer = transformers.BertTokenizer.from_pretrained('bert-base-uncased')
            elif MODEL == 'DistilBERT':
                tokenizer = transformers.DistilBertTokenizer.from_pretrained('distilbert-base-uncased')
            else:
                continue

                
        if LANG == 'FR':
            tfidf_vectorizer = TfidfVectorizer(smooth_idf=True, use_idf=True,
                                           stop_words=list(fr_stop),
                                           tokenizer=tokenizer.tokenize)
        else:
            tfidf_vectorizer = TfidfVectorizer(smooth_idf=True, use_idf=True,
                                           stop_words=list(en_stop),
                                           tokenizer=tokenizer.tokenize)


        # just send in all your docs here
        fitted_vectorizer = tfidf_vectorizer.fit(tdrs)
        tfidf_tdrs = fitted_vectorizer.transform(tdr)

        df_tdrs = pd.DataFrame(tfidf_tdrs.T.todense())
        tfidf_cvs = fitted_vectorizer.transform(cvs)
        df_cvs = pd.DataFrame(tfidf_cvs.T.todense())

        # cosinus score
        actual_score_cos = df_cvs.T.dot(df_tdrs)

        actual_score_cos.columns = ['tdr%d'%i for i in range(len(tdr))]
        actual_score_cos.index = cvs_names
        score_cos += actual_score_cos

        print(strftime("%Y-%m-%d %H:%M:%S", gmtime()))
    score_cos=100*score_cos/3 if LANG=='FR' else 100*score_cos/2

    df_result = score_cos.sort_values(by='tdr0', ascending=False)
    df_result=df_result.drop_duplicates()

    res = df_result.to_json(orient='columns')
    if output is not None:
        df_result.to_excel(output)
    print(strftime("%Y-%m-%d %H:%M:%S", gmtime()))

    return res


# MAIN PART THAT CALLS ALL OTHER FUNCTIONS

import gradio as gr
import docx2txt
from pdfminer.high_level import extract_text
import os
from langdetect import detect

def get_text_from_file(file_path):
    if file_path.endswith('.pdf'):
        return extract_text(file_path)
    elif file_path.endswith('.docx'):
        return docx2txt.process(file_path)
    else:
        raise ValueError("Unsupported file format")

def detect_language(text):
    try:
        lang = detect(text)
        return lang
    except:
        return "unknown"

def categorize_cvs(cv_files):
    english_cvs = []
    french_cvs = []

    for cv_file in cv_files:
        cv_content = get_text_from_file(cv_file)

        # Assuming you have a language detection library (like langdetect) installed
        language = detect(cv_content)

        if language == 'en':
            english_cvs.append(cv_file)
        elif language == 'fr':
            french_cvs.append(cv_file)

    return english_cvs, french_cvs

def generate_json_data(result, cv_files):
    data = []

    for entry in result:
        for filename, score in entry.items():
            language = filename.split("_")[-1].split(".")[0] if "_" in filename else "en"
            full_path = [file_path for file_path in cv_files if os.path.basename(file_path) == filename]

            if full_path:
                file_data = {
                    "full_path": os.path.abspath(full_path[0]),
                    "language": language,
                    "score": score
                }
            else:
                file_data = {
                    "full_path": None,  # You may want to handle the case when full path is not available
                    "language": language,
                    "score": score
                }

            data.append(file_data)

    return data


import concurrent.futures
import json

def compare_job_descriptions_and_cvs(english_job_description, french_job_description, cv_files):
    with concurrent.futures.ThreadPoolExecutor() as executor:
        # Convert job descriptions to JSON concurrently
        future_english_json = executor.submit(convert_file_to_json, english_job_description)
        future_french_json = executor.submit(convert_file_to_json, french_job_description, fr=True)

        # Categorize CVs concurrently
        future_categorized_cvs = executor.submit(categorize_cvs, cv_files)
        english_cvs, french_cvs = future_categorized_cvs.result()

        # Convert CVs to JSON concurrently
        future_english_cv_json = executor.submit(convert_paths_to_json, english_cvs)
        future_french_cv_json = executor.submit(convert_paths_to_json, french_cvs, fr=True)

        # Wait for CVs to JSON conversion tasks to complete for both languages
        concurrent.futures.wait([future_english_json, future_french_json])
        concurrent.futures.wait([future_english_cv_json, future_french_cv_json])

        # Perform ranking concurrently after JSON conversion for each language
        future_english_rank = executor.submit(
            ranking, "./CVs.json", "./CVs.json", "./tdr.json", stop=True, output='./ranking_results.xlsx'
        )
        future_french_rank = executor.submit(
            ranking, "./CVs_fr.json", "./CVs_fr.json", "./tdr_fr.json", stop=True, output='./ranking_results.xlsx'
        )

        # Wait for ranking tasks to complete for both languages
        concurrent.futures.wait([future_english_rank, future_french_rank])

    # Retrieve results
    english_json = future_english_json.result()
    french_json = future_french_json.result()

    english_rank = future_english_rank.result()
    french_rank = future_french_rank.result()

    # Process results
    results = [english_rank, french_rank]
    results = [json.loads(entry)["tdr0"] for entry in results]
    results = generate_json_data(results, cv_files)

    return results
